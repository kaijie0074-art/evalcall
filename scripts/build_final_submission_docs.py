from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"

NAVY = "142039"
BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
GOLD = "D9A900"
GOLD_SOFT = "FFF8D8"
MUTED = "687386"
LINE = "D9E0E9"
TABLE_FILL = "E8EEF5"
CODE_FILL = "F2F4F7"


OUTPUTS = [
    {
        "source": DOCS / "EvalCall答辩演示脚本-10分钟-20260714.md",
        "target": DOCS / "EvalCall答辩演示脚本-10分钟-20260714.docx",
        "preset": "compact_reference_guide",
        "kicker": "EVALCALL · 决赛路演",
        "subtitle": "10 分钟 PPT + 现场 Demo · 计划 9 分 30 秒，预留 30 秒机动",
        "header": "EvalCall 决赛路演讲稿",
    },
    {
        "source": DOCS / "EvalCall答辩问答预案-10分钟-20260714.md",
        "target": DOCS / "EvalCall答辩问答预案-10分钟-20260714.docx",
        "preset": "compact_reference_guide",
        "kicker": "EVALCALL · 评委问答",
        "subtitle": "先结论 · 再证据 · 最后边界 · 单题 30–60 秒",
        "header": "EvalCall 决赛答辩问答预案",
    },
    {
        "source": DOCS / "EvalCall方案设计文档-最终版-20260714.md",
        "target": DOCS / "EvalCall方案设计文档-最终版-20260714.docx",
        "preset": "standard_business_brief",
        "kicker": "EVALCALL · SOLUTION DESIGN",
        "subtitle": "自动测试并评估外呼模型在特定任务指令下的指令遵循效果",
        "header": "EvalCall 方案设计文档 · 决赛最终版",
    },
    {
        "source": DOCS / "EvalCall决赛提交清单-20260714.md",
        "target": DOCS / "EvalCall决赛提交清单-20260714.docx",
        "preset": "compact_reference_guide",
        "kicker": "EVALCALL · FINAL CHECKLIST",
        "subtitle": "2026-07-15 提交 · 10 分钟路演 + 10 分钟答辩",
        "header": "EvalCall 决赛提交与现场清单",
    },
]


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, *, size=None, bold=None, color=None, italic=None, mono=False):
    family = "Menlo" if mono else "Arial Unicode MS"
    east_asia = "Arial Unicode MS"
    run.font.name = family
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), family)
    r_pr.rFonts.set(qn("w:hAnsi"), family)
    r_pr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = rgb(color)
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cell_borders(cell, color=LINE, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def keep_row_together(row, repeat=False):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    if repeat:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def add_field_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc: Document, preset: str):
    body_line = 1.25 if preset == "compact_reference_guide" else 1.10
    body_after = 6
    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(NAVY)
    normal.paragraph_format.space_after = Pt(body_after)
    normal.paragraph_format.line_spacing = body_line

    heading_tokens = {
        "Heading 1": (16, BLUE, 18 if preset == "compact_reference_guide" else 16, 10 if preset == "compact_reference_guide" else 8),
        "Heading 2": (13, BLUE, 14 if preset == "compact_reference_guide" else 12, 7 if preset == "compact_reference_guide" else 6),
        "Heading 3": (12, DEEP_BLUE, 10 if preset == "compact_reference_guide" else 8, 5 if preset == "compact_reference_guide" else 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(11)
        style.font.color.rgb = rgb(NAVY)
        style.paragraph_format.left_indent = Inches(0.375 if preset == "compact_reference_guide" else 0.5)
        style.paragraph_format.first_line_indent = Inches(-0.188 if preset == "compact_reference_guide" else -0.25)
        style.paragraph_format.space_after = Pt(4 if preset == "compact_reference_guide" else 8)
        style.paragraph_format.line_spacing = body_line


def prepare_document(title: str, subtitle: str, kicker: str, header_text: str, preset: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc, preset)

    header = section.header.paragraphs[0]
    header.text = header_text
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        set_run_font(run, size=8.5, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    add_field_page_number(footer)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(kicker)
    set_run_font(run, size=9, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_run_font(run, size=25, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(subtitle)
    set_run_font(run, size=12, color=MUTED)

    callout = doc.add_table(rows=1, cols=2)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout.cell(0, 0).text = "FINAL"
    callout.cell(0, 1).text = "决赛提交版 · 2026-07-14 · 所有数字以仓库正式运行产物为准"
    set_table_geometry(callout, [1500, 7860])
    for index, cell in enumerate(callout.rows[0].cells):
        shade_cell(cell, GOLD_SOFT)
        cell_borders(cell, GOLD, "8")
        cell_margins(cell, 120, 150, 120, 150)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9.5 if index == 0 else 10.5, bold=index == 0, color=NAVY)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return doc


def add_inline_runs(paragraph, text: str, *, base_size=11, base_color=NAVY, mono=False):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=base_size, color=base_color, mono=mono)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=base_color, bold=True, mono=mono)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size - 0.5, color=DEEP_BLUE, mono=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, color=base_color, mono=mono)


def add_paragraph(doc: Document, text: str, preset: str, style=None, italic=False):
    p = doc.add_paragraph(style=style)
    add_inline_runs(p, text)
    p.paragraph_format.space_after = Pt(4 if style in ("List Bullet", "List Number") else 6)
    p.paragraph_format.line_spacing = 1.25 if preset == "compact_reference_guide" else 1.10
    if italic:
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = rgb(MUTED)
    return p


def next_numbering_id(elements, tag: str, attr: str) -> int:
    values = []
    for element in elements.findall(qn(tag)):
        raw = element.get(qn(attr))
        if raw is not None and raw.isdigit():
            values.append(int(raw))
    return max(values, default=0) + 1


def create_decimal_numbering(doc: Document, preset: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = next_numbering_id(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540" if preset == "compact_reference_guide" else "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540" if preset == "compact_reference_guide" else "720")
    ind.set(qn("w:hanging"), "270" if preset == "compact_reference_guide" else "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_paragraph(doc: Document, text: str, preset: str, num_id: int):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    p.paragraph_format.left_indent = Inches(0.375 if preset == "compact_reference_guide" else 0.5)
    p.paragraph_format.first_line_indent = Inches(-0.188 if preset == "compact_reference_guide" else -0.25)
    p.paragraph_format.space_after = Pt(4 if preset == "compact_reference_guide" else 8)
    p.paragraph_format.line_spacing = 1.25 if preset == "compact_reference_guide" else 1.10
    add_inline_runs(p, text)
    return p


def add_callout(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, GOLD_SOFT)
    cell_borders(cell, GOLD, "8")
    cell_margins(cell, 130, 170, 130, 170)
    cell.text = ""
    add_inline_runs(cell.paragraphs[0], text, base_size=10.5)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc: Document, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, CODE_FILL)
    cell_borders(cell, LINE, "6")
    cell_margins(cell, 130, 170, 130, 170)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, size=8.5, color=NAVY, mono=True)
        if index < len(lines) - 1:
            run.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown_table(doc: Document, header: list[str], rows: list[list[str]]):
    cols = len(header)
    if cols == 2:
        widths = [2400, 6960]
    elif cols == 3:
        widths = [1900, 3500, 3960]
    elif cols == 4:
        widths = [1200, 2700, 2300, 3160]
    else:
        base = 9360 // cols
        widths = [base] * cols
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, value in enumerate(header):
        cell = table.cell(0, idx)
        cell.text = ""
        add_inline_runs(cell.paragraphs[0], value, base_size=9.2)
        shade_cell(cell, TABLE_FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_borders(cell)
        cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = rgb(DEEP_BLUE)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    keep_row_together(table.rows[0], repeat=True)

    for row_values in rows:
        row = table.add_row()
        keep_row_together(row)
        values = row_values + [""] * (cols - len(row_values))
        for idx, value in enumerate(values[:cols]):
            cell = row.cells[idx]
            cell.text = ""
            add_inline_runs(cell.paragraphs[0], value, base_size=9.0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_borders(cell)
            cell_margins(cell)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.12
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def split_table_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def render_markdown(doc: Document, markdown: str, preset: str):
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        if not line.strip():
            index += 1
            continue
        if index == 0 and line.startswith("# "):
            index += 1
            continue
        if line.startswith("```"):
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(doc, code_lines)
            index += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline_runs(p, line[4:], base_size=12, base_color=DEEP_BLUE)
            index += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_runs(p, line[3:], base_size=13, base_color=BLUE)
            index += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline_runs(p, line[2:], base_size=16, base_color=BLUE)
            index += 1
            continue
        if line.startswith("> "):
            quote_lines = []
            while index < len(lines) and lines[index].startswith(">"):
                quote_lines.append(lines[index].lstrip("> "))
                index += 1
            add_callout(doc, " ".join(quote_lines))
            continue
        if "|" in line and index + 1 < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|", lines[index + 1]):
            header = split_table_row(line)
            index += 2
            rows = []
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                rows.append(split_table_row(lines[index]))
                index += 1
            add_markdown_table(doc, header, rows)
            continue
        if re.match(r"^\s*[-*] \[[ xX]\] ", line):
            text = re.sub(r"^\s*[-*] \[[ xX]\] ", "☐ ", line)
            add_paragraph(doc, text, preset, style="List Bullet")
            index += 1
            continue
        if re.match(r"^\s*[-*] ", line):
            text = re.sub(r"^\s*[-*] ", "", line)
            add_paragraph(doc, text, preset, style="List Bullet")
            index += 1
            continue
        if re.match(r"^\s*\d+\. ", line):
            num_id = create_decimal_numbering(doc, preset)
            while index < len(lines) and re.match(r"^\s*\d+\. ", lines[index].rstrip()):
                text = re.sub(r"^\s*\d+\. ", "", lines[index].rstrip())
                add_numbered_paragraph(doc, text, preset, num_id)
                index += 1
            continue

        para_lines = [line.strip()]
        index += 1
        while index < len(lines):
            nxt = lines[index].rstrip()
            if not nxt.strip() or nxt.startswith(("#", ">", "```")):
                break
            if re.match(r"^\s*([-*] |\d+\. )", nxt):
                break
            if "|" in nxt and index + 1 < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|", lines[index + 1]):
                break
            para_lines.append(nxt.strip())
            index += 1
        add_paragraph(doc, " ".join(para_lines), preset)


def build_one(config: dict):
    markdown = config["source"].read_text(encoding="utf-8")
    title = markdown.splitlines()[0].removeprefix("# ").strip()
    doc = prepare_document(
        title=title,
        subtitle=config["subtitle"],
        kicker=config["kicker"],
        header_text=config["header"],
        preset=config["preset"],
    )
    render_markdown(doc, markdown, config["preset"])
    doc.core_properties.title = title
    doc.core_properties.subject = "EvalCall 美团黑客松决赛提交物料"
    doc.core_properties.author = "EvalCall Team"
    doc.core_properties.keywords = "EvalCall, 外呼模型, 指令遵循, 美团黑客松"
    doc.save(config["target"])
    return config["target"]


def main():
    outputs = [str(build_one(config)) for config in OUTPUTS]
    print("\n".join(outputs))


if __name__ == "__main__":
    main()
